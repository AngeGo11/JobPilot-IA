#!/usr/bin/env python3
# ---------------------------------------------------------------------------
# Dépendance requise : pip install python-docx
# Génère docs/JobPilot_Guide_Utilisateur.docx
# Exécution : python3 scripts/build_guide_utilisation_docx.py
# ---------------------------------------------------------------------------
"""
Génère le Guide Utilisateur Officiel JobPilot au format Word (.docx).
Utilise la librairie python-docx pour la structure, titres, tableaux et mise en forme.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT_FILE = DOCS_DIR / "JobPilot_Guide_Utilisateur.docx"


def set_document_style(doc):
    """Applique une police et taille par défaut au style Normal."""
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.name = "Calibri"


def add_para_with_bold(doc, text, bold_terms=None):
    """
    Ajoute un paragraphe en mettant en gras les termes de bold_terms.
    bold_terms : liste de chaînes à mettre en gras (ex: ["Crédits", "Premium"]).
    Toutes les occurrences de chaque terme sont mises en gras.
    """
    if not bold_terms:
        return doc.add_paragraph(text)
    p = doc.add_paragraph()
    remaining = text
    while remaining:
        # Trouver la première occurrence de n'importe quel terme
        first_pos = len(remaining)
        first_term = None
        for term in bold_terms:
            idx = remaining.find(term)
            if idx != -1 and idx < first_pos:
                first_pos = idx
                first_term = term
        if first_term is None:
            p.add_run(remaining)
            break
        if first_pos > 0:
            p.add_run(remaining[:first_pos])
        run = p.add_run(first_term)
        run.bold = True
        remaining = remaining[first_pos + len(first_term):]
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def build_guide():
    doc = Document()
    set_document_style(doc)

    # ---------- Titre du document (Style Title) ----------
    title_para = doc.add_paragraph("Guide Utilisateur Officiel - JobPilot")
    title_para.style = doc.styles["Title"]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(24)
    doc.add_paragraph()
    doc.add_page_break()

    # ==================== INTRODUCTION ====================
    add_heading(doc, "Introduction", 1)
    add_para_with_bold(
        doc,
        "Bienvenue sur JobPilot. Cette plateforme vous accompagne dans votre recherche d'emploi en analysant vos CV, "
        "en trouvant des offres pertinentes sur France Travail et en vous aidant à rédiger des lettres de motivation personnalisées.",
        ["JobPilot"],
    )
    doc.add_paragraph()
    add_heading(doc, "Les 3 crédits offerts à l'inscription", 2)
    add_para_with_bold(
        doc,
        "Dès votre inscription, vous recevez 3 Crédits gratuits (ou 5 selon la configuration en place). Ces Crédits permettent "
        "d'utiliser les fonctionnalités pilotées par l'intelligence artificielle : analyse de votre CV (détection du poste visé et des compétences), "
        "génération ou amélioration de lettres de motivation, et optimisation de votre CV pour une offre. Une fois ces Crédits épuisés, "
        "vous pouvez acheter un Pack Recharge (10 crédits) ou souscrire un abonnement Premium pour un accès illimité pendant la durée choisie.",
        ["Crédits", "Crédits", "Crédits", "Pack Recharge", "Premium"],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ==================== COMPRENDRE L'INTERFACE ====================
    add_heading(doc, "Comprendre l'interface", 1)

    add_heading(doc, "Le Dashboard", 2)
    add_para_with_bold(
        doc,
        "Le Dashboard est votre tableau de bord principal. En haut de la page, vous retrouvez des indicateurs clés : "
        "Total des candidatures, Nouveaux matches, Vus, Postulés. Ces indicateurs vous permettent de suivre en un coup d'œil "
        "l'état de votre recherche. La liste des candidatures s'affiche en dessous ; en cliquant sur une ligne, vous ouvrez "
        "le workspace de la candidature (offre + éditeur de lettre de motivation).",
        ["Dashboard"],
    )
    doc.add_paragraph()

    add_heading(doc, "Compte gratuit vs badge « Illimité » (Premium)", 2)
    add_para_with_bold(
        doc,
        "En haut à droite de l'écran, la barre de navigation affiche soit votre solde de Crédits (ex. « 3 Crédits ») si vous êtes "
        "en compte gratuit, soit un badge « Abonné Premium » (ou « Illimité ») avec la date de fin d'abonnement. En compte gratuit : "
        "chaque action IA (analyse CV, génération de lettre, amélioration, optimisation CV) consomme 1 Crédit. Avec un abonnement Premium : "
        "ces actions sont illimitées pendant la période d'abonnement, et vous avez en plus accès aux alertes e-mail pour les nouvelles offres.",
        ["Crédits", "Crédits", "Premium", "Premium", "Crédit"],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ==================== FONCTIONNALITÉS CLÉS (TUTORIELS) ====================
    add_heading(doc, "Fonctionnalités clés (tutoriels)", 1)

    add_heading(doc, "Matching : comment lire le score", 2)
    add_para_with_bold(
        doc,
        "Lorsque vous lancez une recherche d'offres à partir d'un CV, chaque offre trouvée reçoit un score de matching (pourcentage, de 0 à 100). "
        "Ce score indique la pertinence entre le contenu de votre CV et la description de l'offre (mots-clés communs, compétences). "
        "Un score élevé (par exemple 70 % ou plus) signifie une bonne adéquation ; vous pouvez prioriser ces offres pour rédiger vos lettres "
        "et postuler. Le score est affiché sur chaque carte de résultat et dans le Dashboard.",
        ["score de matching", "score"],
    )
    doc.add_paragraph()

    add_heading(doc, "Générateur de lettre de motivation : coût (1 crédit) et personnalisation", 2)
    add_para_with_bold(
        doc,
        "Dans le workspace d'une candidature, le bouton « Générer avec l'IA » (ou équivalent) crée une lettre de motivation personnalisée "
        "à partir de votre CV et de l'offre. Cette action coûte 1 Crédit (ou est illimitée en Premium). La lettre est adaptée au poste "
        "et à l'entreprise ; vous pouvez ensuite la modifier dans l'éditeur (ton, longueur, détails) et utiliser les options d'amélioration "
        "(« Rendre plus formel », « Corriger la grammaire », etc., chaque option consommant 1 Crédit si vous n'êtes pas Premium). "
        "Pensez à cliquer sur « Sauvegarder » pour enregistrer la lettre après génération ou modification.",
        ["Générer avec l'IA", "1 Crédit", "Premium", "Crédit", "Sauvegarder"],
    )
    doc.add_paragraph()

    add_heading(doc, "Optimisation CV : adapter son CV à une offre", 2)
    add_para_with_bold(
        doc,
        "Le bouton « Adapter mon CV à cette offre » (dans le workspace) lance une analyse IA qui compare votre CV à l'offre et vous propose : "
        "mots-clés manquants à intégrer, suggestion de résumé ou d'accroche, pistes pour mettre en avant vos expériences. "
        "Cette fonctionnalité consomme 1 Crédit (ou est incluse dans l'abonnement Premium). Les suggestions s'affichent à l'écran ; "
        "vous pouvez les utiliser pour modifier votre CV en dehors de JobPilot (puis ré-uploader le CV si besoin).",
        ["Adapter mon CV à cette offre", "1 Crédit", "Premium"],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ==================== MODÈLE ÉCONOMIQUE ====================
    add_heading(doc, "Modèle économique (section cruciale)", 1)

    add_heading(doc, "Crédits (besoin ponctuel) vs Abonnement (besoin intensif / alertes)", 2)
    add_para_with_bold(
        doc,
        "JobPilot propose deux façons d'accéder aux fonctionnalités IA : les Crédits pour un usage ponctuel, et l'Abonnement Premium "
        "pour un usage intensif et les alertes e-mail.",
        ["Crédits", "Abonnement Premium", "Premium"],
    )
    doc.add_paragraph()

    # Tableau comparatif
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Critère", "Crédits (Pack Recharge)", "Abonnement (Pass 24h / Sprint / Pro)"]
    for col, header_text in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    data = [
        ["Usage", "Ponctuel : chaque action IA = 1 crédit", "Illimité pendant la durée (24h, 7 j ou 30 j)"],
        ["Alertes e-mail", "Non disponibles", "Incluses (nouvelles offres)"],
        ["Idéal pour", "Quelques candidatures ciblées", "Recherche active, veille continue"],
        ["Exemple", "Pack 10 crédits = 4,99 €", "Sprint 7 j = 5,99 € ; Pro 30 j = 14,99 €"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_text
    doc.add_paragraph()

    add_heading(doc, "Fonctionnalités verrouillées (cadenas)", 2)
    add_para_with_bold(
        doc,
        "Certaines fonctionnalités sont réservées aux abonnés Premium et affichent un cadenas ou un message explicite : "
        "les alertes e-mail (notification de nouvelles offres correspondant à votre CV) sont verrouillées en compte gratuit. "
        "En cliquant sur une fonctionnalité verrouillée, vous êtes invité à consulter la page Tarifs pour souscrire un abonnement "
        "ou acheter des Crédits selon vos besoins.",
        ["Premium", "Crédits"],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ==================== FAQ ====================
    add_heading(doc, "FAQ – Questions courantes", 1)

    add_heading(doc, "Les crédits sont-ils perdus si je ne les utilise pas ?", 2)
    add_para_with_bold(
        doc,
        "Les Crédits restent sur votre compte tant que vous ne les utilisez pas. Ils ne « périment » pas. En revanche, "
        "dès qu'une action IA est effectuée (analyse CV, génération ou amélioration de lettre, optimisation CV), 1 Crédit "
        "est débité. Si vous souscrivez un abonnement Premium, vos Crédits restants sont conservés et utilisés après la fin "
        "de l'abonnement si vous ne renouvelez pas.",
        ["Crédits", "Crédits", "Crédit", "Premium", "Crédits"],
    )
    doc.add_paragraph()

    add_heading(doc, "Comment gérer ou annuler mon abonnement ?", 2)
    add_para_with_bold(
        doc,
        "Les abonnements Sprint et Pro se renouvellent automatiquement. Pour annuler ou modifier votre abonnement, vous devez "
        "passer par le tableau de bord Stripe (lien fourni dans l’e-mail de confirmation de paiement) ou contacter le support JobPilot. "
        "Une fois l'abonnement annulé, vous conservez l'accès Premium jusqu'à la date de fin déjà payée ; ensuite, vous repassez "
        "en mode Crédits (votre solde restant) ou devrez acheter un nouveau pack ou abonnement.",
        ["Stripe", "JobPilot", "Premium", "Crédits"],
    )
    doc.add_paragraph()

    add_heading(doc, "Que se passe-t-il quand je n'ai plus de crédits ?", 2)
    add_para_with_bold(
        doc,
        "Si votre solde de Crédits est à 0 et que vous n'êtes pas abonné Premium, les actions IA (génération de lettre, "
        "amélioration, optimisation CV, analyse IA d'un nouveau CV) afficheront un message vous invitant à recharger. "
        "Rendez-vous sur la page « Tarifs » pour acheter un Pack Recharge (10 crédits) ou souscrire un Pass 24h, Sprint ou Pro.",
        ["Crédits", "Premium", "Crédits", "Tarifs", "Pack Recharge"],
    )
    doc.add_paragraph()

    add_heading(doc, "L'export PDF de ma lettre consomme-t-il un crédit ?", 2)
    add_para_with_bold(
        doc,
        "Non. L'export de votre lettre de motivation en PDF ne consomme pas de Crédit. Seules les actions qui font appel "
        "à l'IA (génération, amélioration, optimisation CV, analyse du CV) consomment des Crédits (ou sont illimitées en Premium).",
        ["Crédit", "Crédits", "Premium"],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ---------- Fin du document ----------
    doc.add_paragraph()
    end_para = doc.add_paragraph("— Fin du Guide Utilisateur Officiel JobPilot —")
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in end_para.runs:
        run.italic = True

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Guide généré : {OUTPUT_FILE}")


if __name__ == "__main__":
    build_guide()
