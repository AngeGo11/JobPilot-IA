#!/usr/bin/env python3
# ---------------------------------------------------------------------------
# Dépendance : pip install python-docx
# Génère docs/JobPilot_Dossier_Technique.docx
# Exécution : python3 scripts/generate_tech_specs.py
# ---------------------------------------------------------------------------
"""
Génère le Dossier Technique & Architecture JobPilot au format Word (.docx).
Documentation technique pour l'équipe de développement.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT_FILE = DOCS_DIR / "JobPilot_Dossier_Technique.docx"


def set_document_style(doc):
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.name = "Calibri"


def add_code_block(doc, code_text):
    """Insère un bloc de code avec police monospace (Courier New)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.bold = False
    # Fond gris léger (optionnel, via XML)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text):
    return doc.add_paragraph(text)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def build_dossier():
    doc = Document()
    set_document_style(doc)

    # ==================== PAGE DE GARDE ====================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run("Dossier Technique & Architecture - JobPilot")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Calibri"

    doc.add_paragraph()
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Plateforme SAAS de recrutement assistée par IA")
    sub_run.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.name = "Calibri"

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date : {date.today().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(11)
    date_run.font.name = "Calibri"

    doc.add_page_break()

    # ==================== SOMMAIRE (TABLE DES MATIÈRES MANUELLE) ====================
    add_heading(doc, "Sommaire", 1)
    sommaire = [
        "1. Vue d'ensemble (Stack technique)",
        "2. Architecture de la base de données (Models)",
        "3. Logique métier critique (Core Business Logic)",
        "4. Sécurité et authentification",
        "5. Frontend et UX",
        "6. Procédure de déploiement (DevOps)",
    ]
    for line in sommaire:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_page_break()

    # ==================== 1. VUE D'ENSEMBLE (STACK TECHNIQUE) ====================
    add_heading(doc, "1. Vue d'ensemble (Stack technique)", 1)

    add_heading(doc, "1.1 Backend", 2)
    add_para(doc, "Django 5.2 (Python 3.11+). Framework web MVT, gestion auth, sessions, ORM, admin.")
    add_bullet(doc, "Applications métier : users, resumes, matching, dashboard, subscriptions.")

    add_heading(doc, "1.2 Frontend", 2)
    add_para(doc, "Django Templates (moteur de templates) + Tailwind CSS. Approche Mobile-First ; pas de framework JS dédié côté front (requêtes AJAX ponctuelles pour les actions IA).")
    add_bullet(doc, "Tailwind chargé via CDN (cdn.tailwindcss.com) ; configuration JIT possible via script tailwind.config dans base.html.")
    add_bullet(doc, "Polices : Outfit, Space Grotesk (Google Fonts). Icônes : Font Awesome 6. Éditeur riche : TinyMCE 8 (lettres de motivation).")

    add_heading(doc, "1.3 Base de données", 2)
    add_para(doc, "PostgreSQL en production et en développement (configuration via .env : DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT). Port par défaut 5433 (Docker). Pas de SQLite en configuration actuelle.")

    add_heading(doc, "1.4 IA", 2)
    add_para(doc, "Google Gemini API (google-generativeai). Modèle utilisé : gemini-2.5-flash. Clé via variable d'environnement GEMINI_API_KEY.")
    add_bullet(doc, "Usages : analyse de CV (titre de poste, compétences), génération et raffinement de lettres de motivation, optimisation CV pour une offre.")

    add_heading(doc, "1.5 Paiement", 2)
    add_para(doc, "Stripe : Checkout Session pour les achats (Pass 24h, Sprint, Pro, Pack Recharge). Webhooks pour checkout.session.completed, customer.subscription.updated, customer.subscription.deleted. Clés : STRIPE_SECRET_KEY, STRIPE_PUBLISHABLE_KEY, STRIPE_WEBHOOK_SECRET ; Price IDs par plan (STRIPE_PRICE_PASS24H, etc.).")

    add_heading(doc, "1.6 Autres services externes", 2)
    add_para(doc, "France Travail API (OAuth2 partenaire) : recherche d'offres (mots-clés, pagination). Variables : ID_CLIENT, CLIENT_SECRET, API_BASE_URL. Extraction PDF : pdfplumber. Export PDF lettres : reportlab.")
    doc.add_page_break()

    # ==================== 2. ARCHITECTURE BASE DE DONNÉES (MODELS) ====================
    add_heading(doc, "2. Architecture de la base de données (Models)", 1)

    add_heading(doc, "2.1 CustomUser (users.models)", 2)
    add_para(doc, "Modèle utilisateur personnalisé (AUTH_USER_MODEL = 'users.CustomUser'), hérite de AbstractUser. Champs critiques pour le modèle économique :")
    add_bullet(doc, "ai_credits (IntegerField) : solde de crédits IA. Défaut 5 à l'inscription. Consommé à chaque action IA (analyse CV, génération/amélioration lettre, optimisation CV) si l'utilisateur n'est pas Premium.")
    add_bullet(doc, "subscription_end_date (DateTimeField, null=True, blank=True) : date/heure de fin d'abonnement. Si dans le futur, l'utilisateur est considéré Premium (accès illimité aux actions IA et alertes).")
    add_bullet(doc, "subscription_plan (CharField, choices=SubscriptionPlan) : pass24h, sprint, pro. Permet d'afficher le type d'abonnement (get_subscription_plan_display).")
    add_para(doc, "Propriétés calculées (pas en BDD) :")
    add_bullet(doc, "is_premium : True si subscription_end_date > now().")
    add_bullet(doc, "can_generate : True si is_premium ou ai_credits > 0.")
    add_code_block(doc, """class CustomUser(AbstractUser):
    ai_credits = models.IntegerField("Crédits IA", default=5)
    subscription_end_date = models.DateTimeField("Fin d'abonnement", null=True, blank=True)
    subscription_plan = models.CharField(max_length=20, choices=SubscriptionPlan.choices, null=True, blank=True)

    @property
    def is_premium(self):
        if not self.subscription_end_date:
            return False
        return self.subscription_end_date > timezone.now()

    @property
    def can_generate(self):
        return self.is_premium or (self.ai_credits and self.ai_credits > 0)""")

    add_heading(doc, "2.2 CandidateProfile (users.models)", 2)
    add_para(doc, "OneToOneField vers CustomUser (related_name='profile'). Champs : phone, location, target_rome_code, is_available. Créé à l'inscription (get_or_create dans la vue register).")

    add_heading(doc, "2.3 Resume (resumes.models)", 2)
    add_para(doc, "ForeignKey vers CustomUser. Représente un CV uploadé (fichier PDF). Champs clés : title, file (upload_to='cvs/'), is_primary ; extracted_text (pdfplumber), detected_job_title, detected_skills (IA), parsed_data (JSON). Un utilisateur a plusieurs Resume ; les offres sont matchées par CV (JobMatch lie resume + job_offer).")

    add_heading(doc, "2.4 JobOffer (matching.models)", 2)
    add_para(doc, "Offres importées depuis France Travail. remote_id (unique), title, company_name, description, url, location, contract_type, date_posted, raw_api_data (JSON). Une même offre peut être liée à plusieurs JobMatch (un par CV candidat).")

    add_heading(doc, "2.5 JobMatch (matching.models) — Candidatures générées", 2)
    add_para(doc, "Table de liaison (résultat du matching) : « Ce CV (Resume) matche avec cette Offre (JobOffer) à X% ». ForeignKey resume (null=True), user, job_offer ; score (IntegerField 0–100), status (new/seen/applied/rejected), matched_at, cover_letter_content. Contrainte unique_together (resume, job_offer). Ordering par score décroissant.")
    add_code_block(doc, """class JobMatch(models.Model):
    resume = models.ForeignKey(Resume, on_delete=models.CASCADE, related_name='matches', null=True)
    user = models.ForeignKey(settings.AUTH_USER_MODEL, on_delete=models.CASCADE)
    job_offer = models.ForeignKey(JobOffer, on_delete=models.CASCADE)
    score = models.IntegerField("Score de matching", default=0)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='new')
    cover_letter_content = models.TextField("Lettre de motivation", blank=True)
    class Meta:
        unique_together = ('resume', 'job_offer')
        ordering = ['-score']""")

    add_heading(doc, "2.6 JobAlert (matching.models)", 2)
    add_para(doc, "ForeignKey Resume, OneToOne par resume (unique_together sur resume). is_active, last_checked. Utilisé par la commande management check_new_offers pour envoyer des e-mails (alertes réservées Premium).")

    add_heading(doc, "2.7 StripeSubscription (subscriptions.models)", 2)
    add_para(doc, "OneToOneField User. stripe_subscription_id (unique). Permet de lier un abonnement Stripe à l'utilisateur et de mettre à jour subscription_end_date lors des webhooks subscription.updated / subscription.deleted.")

    add_heading(doc, "2.8 Schéma relationnel résumé", 2)
    add_para(doc, "User 1 — 1 CandidateProfile ; User 1 — N Resume ; Resume 1 — N JobMatch ; JobOffer 1 — N JobMatch ; (resume, job_offer) unique pour JobMatch. User 1 — 1 StripeSubscription. Resume 1 — 1 JobAlert.")
    doc.add_page_break()

    # ==================== 3. LOGIQUE MÉTIER CRITIQUE ====================
    add_heading(doc, "3. Logique métier critique (Core Business Logic)", 1)

    add_heading(doc, "3.1 Système de crédits (transactionnel)", 2)
    add_para(doc, "La méthode consume_credit(user) (matching.services.credits) est le point central de débit des crédits. Elle doit être appelée avant toute action IA (génération lettre, amélioration, optimisation CV, analyse IA du CV) pour les utilisateurs non Premium.")
    add_para(doc, "Comportement : si user.is_premium, retourne True sans déduire. Sinon, mise à jour atomique en base : filter(pk=user.pk, ai_credits__gt=0).update(ai_credits=F('ai_credits') - 1). F() évite les race conditions en effectuant le décrément côté base (read-modify-write atomique). Si update retourne 1, user.refresh_from_db() et return True ; sinon return False.")
    add_code_block(doc, """def consume_credit(user):
    if user.is_premium:
        return True
    updated = user.__class__.objects.filter(
        pk=user.pk,
        ai_credits__gt=0
    ).update(ai_credits=F('ai_credits') - 1)
    if updated:
        user.refresh_from_db()
        return True
    return False""")
    add_para(doc, "Renforcement possible contre les race conditions (double dépense en concurrence forte) : envelopper l'appel dans transaction.atomic() et utiliser select_for_update() sur l'utilisateur avant de vérifier ai_credits et de décrémenter (avec une mise à jour par F() dans la même transaction). L'implémentation actuelle avec F() seule est déjà sûre pour la majorité des cas.")

    add_heading(doc, "3.2 Feature locking (Freemium)", 2)
    add_para(doc, "Distinction entre accès par Crédits (consommable : chaque action IA = 1 crédit) et Premium (illimité pendant subscription_end_date).")
    add_bullet(doc, "Vérification côté vue : avant d'exécuter une action IA, appel à consume_credit(request.user) ; si False, renvoi JsonResponse 402 ou redirect vers page Tarifs avec message « Crédits insuffisants ».")
    add_bullet(doc, "Helpers sur le modèle : user.is_premium (property) et user.can_generate (property). Pas de décorateur @premium_required dans le code actuel ; les vues qui réservent une fonctionnalité au Premium (ex. toggle_job_alert) testent if not request.user.is_premium puis retournent 403.")
    add_para(doc, "Exemple de pattern à réutiliser pour une vue « Premium only » :")
    add_code_block(doc, """if not getattr(request.user, 'is_premium', False):
    return JsonResponse({
        'success': False,
        'error': "Les alertes email sont réservées aux abonnés Premium.",
        'redirect': '/subscriptions/pricing/',
    }, status=403)""")
    add_para(doc, "Un décorateur @premium_required pourrait être ajouté (décorator qui vérifie request.user.is_premium et redirige ou renvoie 403) pour centraliser la logique sur les vues réservées aux abonnés.")
    doc.add_page_break()

    # ==================== 4. SÉCURITÉ ET AUTHENTIFICATION ====================
    add_heading(doc, "4. Sécurité et authentification", 1)

    add_heading(doc, "4.1 django-allauth (Social Login)", 2)
    add_para(doc, "Authentification sociale : Google, GitHub, LinkedIn (allauth.socialaccount.providers.google / github / linkedin_oauth2). SOCIALACCOUNT_LOGIN_ON_GET = True. Routes sous /accounts/ (include('allauth.urls')). Connexion possible par formulaire classique (username/password) ou via fournisseur OAuth.")

    add_heading(doc, "4.2 Variables d'environnement (.env)", 2)
    add_para(doc, "Toutes les clés sensibles sont chargées depuis l'environnement (python-dotenv, load_dotenv() en tête de settings.py). À ne jamais commiter : SECRET_KEY, DB_PASSWORD, ID_CLIENT, CLIENT_SECRET, GEMINI_API_KEY, STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET. ALLOWED_HOSTS et CSRF_TRUSTED_ORIGINS doivent être configurés pour le domaine de production.")

    add_heading(doc, "4.3 Authentification et autorisation", 2)
    add_bullet(doc, "Login requis : @login_required sur les vues métier (dashboard, matching, resumes, workspace). LOGIN_URL = 'login', LOGIN_REDIRECT_URL = 'post_login_loading'.")
    add_bullet(doc, "Vérification de propriété : get_object_or_404(JobMatch, id=match_id, user=request.user) pour garantir que l'utilisateur n'accède qu'à ses propres candidatures.")
    add_bullet(doc, "CSRF : CsrfViewMiddleware actif ; formulaire avec {% csrf_token %}. Webhook Stripe : @csrf_exempt (obligatoire) avec vérification de signature via stripe.Webhook.construct_event(payload, sig_header, webhook_secret).")
    doc.add_page_break()

    # ==================== 5. FRONTEND ET UX ====================
    add_heading(doc, "5. Frontend et UX", 1)

    add_heading(doc, "5.1 Configuration Tailwind", 2)
    add_para(doc, "Tailwind chargé via CDN (script src=\"https://cdn.tailwindcss.com\"). Configuration JIT dans un bloc <script> tailwind.config : extend theme (fontFamily : sans = Outfit, display = Space Grotesk). Pas de build Node en place ; pour la production, un build Tailwind (npm + purge) peut être ajouté pour réduire la taille CSS.")

    add_heading(doc, "5.2 Stratégie responsive", 2)
    add_para(doc, "Points de rupture : md: et lg: (breakpoints Tailwind). Navigation : menu horizontal en md+, menu hamburger et panneau déroulant en dessous de md. Grilles : grid-cols-1 md:grid-cols-2 lg:grid-cols-3 pour les cartes (offres, tarifs). Workspace candidature : disposition split (offre / lettre) adaptée en colonne sur petit écran.")

    add_heading(doc, "5.3 Composants UI et fonctionnalités verrouillées", 2)
    add_para(doc, "Fonctionnalités réservées Premium (ex. alertes e-mail) sont signalées par un message explicite et un lien vers la page Tarifs. Pattern « Glassmorphism » ou effet visuel de verrouillage : cartes ou boutons avec style atténué (opacité, bordure) et icône cadenas ou texte « Réservé aux abonnés Premium » ; au clic, redirection vers /subscriptions/pricing/ (ou équivalent).")
    doc.add_page_break()

    # ==================== 6. PROCÉDURE DE DÉPLOIEMENT (DEVOPS) ====================
    add_heading(doc, "6. Procédure de déploiement (DevOps)", 1)

    add_heading(doc, "6.1 Prérequis", 2)
    add_bullet(doc, "Python 3.11+, PostgreSQL 15+, pip.")
    add_bullet(doc, "Installation des dépendances : pip install -r requirements.txt (ou dans un environnement virtuel recommandé).")
    add_bullet(doc, "Fichier .env à la racine avec toutes les variables nécessaires (DB_*, SECRET_KEY, GEMINI_API_KEY, Stripe, France Travail, etc.).")

    add_heading(doc, "6.2 Base de données", 2)
    add_bullet(doc, "Créer la base PostgreSQL (ou utiliser Docker : docker-compose up -d db).")
    add_bullet(doc, "Commandes : python manage.py migrate (appliquer les migrations).")

    add_heading(doc, "6.3 Serveur d'application (recommandé)", 2)
    add_para(doc, "Gunicorn + Nginx en production. Exemple : gunicorn JobPilot.wsgi:application --bind 0.0.0.0:8000 --workers 4. Nginx en reverse proxy (proxy_pass vers le socket ou le port Gunicorn), gestion des fichiers statiques et optionnellement des médias.")

    add_heading(doc, "6.4 Commandes clés", 2)
    add_code_block(doc, "python manage.py migrate          # Appliquer les migrations\npython manage.py collectstatic  # Collecter les fichiers statiques (DEBUG=False)\npython manage.py createsuperuser # Créer un compte admin")
    add_para(doc, "En production : DEBUG=False, SECRET_KEY forte, ALLOWED_HOSTS et CSRF_TRUSTED_ORIGINS corrects. Configurer un vrai EMAIL_BACKEND pour les e-mails (alertes, réinitialisation mot de passe).")

    add_heading(doc, "6.5 Webhooks Stripe", 2)
    add_para(doc, "Exposer une URL publique pour le webhook (ex. /subscriptions/webhook/) et la configurer dans le tableau de bord Stripe. Utiliser STRIPE_WEBHOOK_SECRET pour la vérification de signature. En local, ngrok ou équivalent pour recevoir les événements.")
    doc.add_paragraph()

    # Fin
    doc.add_paragraph()
    end_para = doc.add_paragraph("— Fin du Dossier Technique JobPilot —")
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in end_para.runs:
        run.italic = True

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Dossier technique généré : {OUTPUT_FILE}")


if __name__ == "__main__":
    build_dossier()
