#!/usr/bin/env python3
"""
Génère docs/ANALYSE_APP_JOBPILOT.docx à partir du contenu de l'analyse.
Exécuter depuis la racine du projet : python scripts/build_analyse_docx.py
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT_FILE = DOCS_DIR / "ANALYSE_APP_JOBPILOT.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    add_heading(doc, "Analyse en profondeur – JobPilot", 0)
    add_paragraph(doc, "Document généré à partir de l’analyse technique de l’application.", bold=False)
    doc.add_paragraph()

    # 1. Vue d'ensemble
    add_heading(doc, "1. Vue d’ensemble", 1)
    add_paragraph(
        doc,
        "JobPilot est une application web Django de recherche d’emploi assistée par IA : "
        "analyse de CV, matching avec des offres (France Travail), génération et amélioration "
        "de lettres de motivation, dashboard de candidatures et monétisation (crédits + abonnements Stripe)."
    )
    doc.add_paragraph()

    # 2. Stack technique
    add_heading(doc, "2. Stack technique", 1)
    p = doc.add_paragraph()
    p.add_run("Backend : ").bold = True
    p.add_run("Django 5.2.10, Python 3.11+")
    add_paragraph(doc, "Base de données : PostgreSQL (port 5433 en dev)")
    add_paragraph(doc, "Auth : django.contrib.auth + django-allauth (Google, GitHub, LinkedIn)")
    add_paragraph(doc, "Frontend : Tailwind CSS, Font Awesome 6, TinyMCE 8, polices Outfit / Space Grotesk")
    add_paragraph(doc, "IA : Google Gemini (gemini-2.5-flash) – analyse CV, lettres, optimisation CV")
    add_paragraph(doc, "Paiement : Stripe (Checkout + webhooks)")
    add_paragraph(doc, "APIs externes : France Travail (OAuth2 partenaire, API offres)")
    add_paragraph(doc, "PDF : pdfplumber (extraction), reportlab (export lettres)")
    add_paragraph(doc, "Infra : Docker Compose (PostgreSQL, Redis, Adminer)")
    doc.add_paragraph()

    # 3. Architecture
    add_heading(doc, "3. Architecture des apps Django", 1)
    add_paragraph(doc, "users : CustomUser (crédits IA, abonnement), CandidateProfile, auth classique + sociale, paramètres.")
    add_paragraph(doc, "resumes : Upload PDF, extraction texte, analyse IA (titre de poste + compétences), liste/suppression, CV principal.")
    add_paragraph(doc, "matching : Recherche offres (France Travail), score de matching, JobMatch / JobAlert, génération/raffinement lettres (IA), export PDF, optimisation CV.")
    add_paragraph(doc, "dashboard : Liste des candidatures, stats (new/seen/applied), workspace détaillé (offre + lettre).")
    add_paragraph(doc, "subscriptions : Page tarifs, création de session Stripe Checkout, succès/annulation, webhook (checkout complété, abo mis à jour/supprimé).")
    doc.add_paragraph()

    # 4. Modèles
    add_heading(doc, "4. Modèles de données (résumé)", 1)
    add_paragraph(doc, "CustomUser : ai_credits (défaut 3), subscription_end_date, subscription_plan (pass24h, sprint, pro). Propriétés : is_premium, can_generate.")
    add_paragraph(doc, "CandidateProfile : 1-1 avec User, téléphone, lieu, code ROME cible, disponibilité.")
    add_paragraph(doc, "Resume : user, fichier PDF, extracted_text, detected_job_title, detected_skills / parsed_skills, parsed_data, is_primary.")
    add_paragraph(doc, "JobOffer : remote_id (France Travail), titre, entreprise, description, url, lieu, type de contrat, raw_api_data.")
    add_paragraph(doc, "JobMatch : resume, user, job_offer, score, status (new/seen/applied/rejected), cover_letter_content ; unique_together (resume, job_offer).")
    add_paragraph(doc, "JobAlert : 1 alerte par CV, is_active, last_checked ; utilisé par la commande check_new_offers.")
    add_paragraph(doc, "StripeSubscription : 1-1 User, stripe_subscription_id pour mettre à jour subscription_end_date via webhooks.")
    doc.add_paragraph()

    # 5. Flux métier
    add_heading(doc, "5. Flux métier principaux", 1)
    add_heading(doc, "5.1 Inscription / Connexion", 2)
    add_paragraph(doc, "Inscription → création user + CandidateProfile → login → session (user_id, email, resume_count, etc.) → page de chargement → redirection dashboard.")
    add_paragraph(doc, "Connexion : même page de chargement puis dashboard.")
    add_paragraph(doc, "Allauth : Google, GitHub, LinkedIn (SOCIALACCOUNT_LOGIN_ON_GET = True).")
    doc.add_paragraph()

    add_heading(doc, "5.2 Parcours CV", 2)
    add_bullet(doc, "Upload PDF → PDFParser (pdfplumber) → extracted_text + parsed_data.")
    add_bullet(doc, "1 crédit consommé (ou gratuit si premium) → AIParser (Gemini) → detected_job_title, detected_skills.")
    add_bullet(doc, "Liste des CVs avec indicateur d’alerte active (premium).")
    doc.add_paragraph()

    add_heading(doc, "5.3 Recherche d’offres", 2)
    add_bullet(doc, "Depuis un CV : page de chargement → find_jobs_for_resume(resume_id).")
    add_bullet(doc, "Recherche France Travail avec resume.detected_job_title (pas de décompte de crédit côté recherche dans le code actuel).")
    add_bullet(doc, "FranceTravail.save_jobs() : création/mise à jour JobOffer, calcul du score (Jaccard sur mots, stopwords français), création JobMatch (resume + user + offre + score).")
    add_bullet(doc, "Résultats paginés (9 par page) pour ce CV uniquement.")
    doc.add_paragraph()

    add_heading(doc, "5.4 Lettres de motivation (workspace)", 2)
    add_bullet(doc, "Sauvegarde : POST classique sur le workspace (contenu + option « marquer postulé »).")
    add_bullet(doc, "Génération (bouton IA) : 1 crédit (ou premium) → AILetterGenerator.generate_cover_letter() → retour JSON (pas de sauvegarde auto).")
    add_bullet(doc, "Amélioration (improve / formalize / grammar / length) : 1 crédit → refine_cover_letter() → sauvegarde du contenu.")
    add_bullet(doc, "Export PDF : pas de crédit → export_to_pdf() (reportlab) → JSON avec PDF en base64 ou téléchargement direct selon la vue.")
    add_bullet(doc, "Optimisation CV (adapter le CV à l’offre) : 1 crédit → AIOptimizer.optimize_for_offer() → suggestions (mots-clés manquants, résumé, expériences).")
    doc.add_paragraph()

    add_heading(doc, "5.5 Crédits et abonnements", 2)
    add_bullet(doc, "Consommation : matching.services.consume_credit(user) — si is_premium : pas de déduction ; sinon décrément atomique de ai_credits (F()).")
    add_bullet(doc, "Stripe : plans pass24h (payment, 1 jour), sprint (subscription, 7 j), pro (subscription, 30 j), pack (payment, 10 crédits). Webhook : checkout.session.completed → apply_plan_to_user (date de fin ou crédits) ; abonnements → StripeSubscription + mise à jour subscription_end_date ; subscription.updated / subscription.deleted pour renouvellement et annulation.")
    doc.add_paragraph()

    # 6. Points forts
    add_heading(doc, "6. Points forts", 1)
    add_bullet(doc, "Séparation claire : apps dédiées, services (France Travail, Gemini, Stripe, crédits) bien isolés.")
    add_bullet(doc, "Modèle utilisateur : crédits + abonnement avec is_premium / can_generate et décompte atomique.")
    add_bullet(doc, "Intégration France Travail : OAuth2, recherche par mots-clés, pagination, sauvegarde des offres et matching par CV.")
    add_bullet(doc, "IA : analyse CV (titre + compétences), lettres (génération + raffinement multi-actions), optimisation CV, tout en Gemini 2.5 Flash.")
    add_bullet(doc, "Monétisation : Stripe Checkout + webhooks pour paiements et abonnements, 4 formules.")
    add_bullet(doc, "Alertes : JobAlert + commande check_new_offers (nouveaux matches + email si score ≥ 70 %, sans spam à la première exécution).")
    add_bullet(doc, "UI : base commune (Tailwind, nav avec crédits/premium), responsive, workspace split-screen pour la candidature.")
    add_bullet(doc, "Logging : config dédiée (fichiers, niveaux, loggers par app).")
    doc.add_paragraph()

    # 7. Points d'attention
    add_heading(doc, "7. Points d’attention et améliorations", 1)
    add_heading(doc, "7.1 Sécurité et robustesse", 2)
    add_bullet(doc, "update_match_status : le commentaire indique une vérification user désactivée ; en production il faut vérifier match.user == request.user pour éviter de modifier le statut d’un autre utilisateur.")
    add_bullet(doc, "Webhook Stripe : @csrf_exempt est normal ; la vérification de signature (construct_event) est bien en place.")
    add_bullet(doc, "SECRET_KEY : dépend de os.getenv(\"SECRET_KEY\") alors que le message d’erreur parle de DJANGO_SECRET_KEY ; à aligner.")
    add_bullet(doc, "Pas de rate limiting explicite sur les vues coûteuses (recherche, IA, Stripe) ; à prévoir pour la prod.")
    doc.add_paragraph()

    add_heading(doc, "7.2 Recherche d’offres et crédits", 2)
    add_bullet(doc, "La recherche d’offres (find_jobs_for_resume) ne consomme pas de crédit. Si le produit prévoit de facturer les recherches, il faut appeler consume_credit (avec règle « 0 résultat = pas de décompte » si souhaité).")
    add_bullet(doc, "Pagination : search_jobs(..., limit=10) alors que la pagination affiche 9 résultats ; à harmoniser.")
    doc.add_paragraph()

    add_heading(doc, "7.3 Données et cohérence", 2)
    add_bullet(doc, "Resume : champs detected_skills et parsed_skills (legacy) ; à terme, un seul champ « skills » et une migration pour fusionner/nettoyer.")
    add_bullet(doc, "Stripe : handle_checkout_completed pour un abonnement sans subscription dans la session (cas limite) pourrait être géré explicitement.")
    doc.add_paragraph()

    add_heading(doc, "7.4 UX et edge cases", 2)
    add_bullet(doc, "Lettre générée : retournée en JSON sans sauvegarde automatique ; l’utilisateur doit cliquer « Sauvegarder ». C’est cohérent avec le code mais à expliquer clairement dans l’UI.")
    add_bullet(doc, "Premium : dans la nav, user.get_subscription_plan_display est utilisé ; Django fournit bien get_subscription_plan_display pour un champ choices, donc correct. Vérifier que SubscriptionPlan couvre bien tous les plans Stripe (pass24h, sprint, pro ; pack = crédits uniquement).")
    doc.add_paragraph()

    add_heading(doc, "7.5 Production", 2)
    add_bullet(doc, "Redis : présent dans docker-compose mais pas utilisé dans settings.py (cache/sessions/Celery) ; soit l’utiliser, soit le retirer du schéma.")
    add_bullet(doc, "Celery : non configuré ; check_new_offers est une commande manuelle/cron. Pour des alertes en temps voulu, une tâche planifiée (Celery Beat ou cron) est nécessaire.")
    add_bullet(doc, "Médias : en DEBUG, servis par Django ; en production, prévoir un service de stockage (S3, etc.) et MEDIA_URL / MEDIA_ROOT adaptés.")
    add_bullet(doc, "Email : EMAIL_BACKEND = console ; en prod, configurer un vrai backend (SMTP, SendGrid, etc.) pour les alertes et les mails système.")
    doc.add_paragraph()

    # 8. Synthèse
    add_heading(doc, "8. Synthèse", 1)
    add_paragraph(doc, "Fonctionnalité cœur : OK – Recherche, matching, lettres IA, dashboard, alertes.")
    add_paragraph(doc, "Monétisation : OK – Crédits + Stripe (4 formules + webhooks).")
    add_paragraph(doc, "Auth : OK – Classique + sociale (Google, GitHub, LinkedIn).")
    add_paragraph(doc, "Sécurité : À renforcer (vérif. user sur update_match_status, rate limiting, RGPD si besoin).")
    add_paragraph(doc, "Recherche / crédits : Clarifier si la recherche doit consommer des crédits et aligner pagination.")
    add_paragraph(doc, "Production : Email, médias, cron/Celery pour alertes, variables d’env et DEBUG.")
    doc.add_paragraph()
    add_paragraph(
        doc,
        "En résumé : architecture solide, fonctionnalités métier et monétisation déjà en place. "
        "Les prochaines étapes logiques sont : sécuriser les vues sensibles, clarifier la politique de crédits pour la recherche, "
        "puis préparer la prod (email, stockage, planification des alertes, monitoring)."
    )

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Fichier généré : {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
